from docx import Document
from pdm_builder.planilhao_build import DocxParser


class DODocBuilder:

    def __init__(self, path_docx_origi, new_file_path):

        self.data = self.parse_data(path_docx_origi)
        self.new_file_path = new_file_path
        self.new_doc = Document()

    def parse_data(self, path_docx):

        parser = DocxParser(path_docx)
        data = parser()

        return data

    def build_obj_estrat(self, meta):

        text = '((NG))Objetivo Estratégico:((CL)) {desc_objetivo}'

        obj = meta['Objetivo estratégico'].strip()

        return text.format(desc_objetivo=obj)

    def build_meta(self, meta):

        text = '((NG))Meta {meta_num}:((CL)) {meta_desc}'

        meta_num = meta['Meta'].strip()
        meta_desc = meta['Meta_'].strip()

        return text.format(meta_num=meta_num, meta_desc=meta_desc)

    def build_indicador(self, meta):

        text = '((NG))Indicador:((CL)) {desc_indi}'

        desc_indi = meta['Indicador'].strip()

        return text.format(desc_indi=desc_indi)

    def build_iniciativas(self, meta):

        text = '((NG))Iniciativas:((CL))\n{iniciativas}'

        inis = meta['iniciativas']
        inis_text = '\n'.join([f'{key} {val}' for key, val in inis.items()])

        return text.format(iniciativas=inis_text)

    def build_page(self, meta, doc):

        obj_estra = self.build_obj_estrat(meta)
        doc.add_paragraph(obj_estra)
        meta_text = self.build_meta(meta)
        doc.add_paragraph(meta_text)
        indicador = self.build_indicador(meta)
        doc.add_paragraph(indicador)
        iniciativas = self.build_iniciativas(meta)
        doc.add_paragraph(iniciativas)

    def build_whole_doc(self, doc=None, data=None):

        if data is None:
            data = self.data
        if doc is None:
            doc = self.new_doc

        for i, meta in enumerate(data):
            self.build_page(meta, doc)
            if i < len(data) - 1:
                doc.add_paragraph('\n')

    def __call__(self):

        self.build_whole_doc()
        self.new_doc.save(self.new_file_path)

if __name__ == '__main__':

    from pathlib import Path

    path_onedrive = Path(r'C:\Users\h-pgy\one_drive_prefs\OneDrive - Default Directory\Shared Documents\Versão final do Programa de Metas')

    path_arquivo = path_onedrive/'Fichas FINAL - 30.06.docx'

    builder = DODocBuilder(path_arquivo, 'word_diario_oficial_final_limpo.docx')
    builder()