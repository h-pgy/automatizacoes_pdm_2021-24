from pathlib import Path
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdm_builder.ficha_parsing import ParserFichas
from pdm_builder.doc_builder import TableBuilder
from pdm_builder.map_builder import MapBuilder
from pdm_builder.tools import pegar_filtro_alteracoes, sort_fichas, checar_metas_presentes, get_map_files

PATH_FILES_ONE_DRIVE = Path(r'C:\Users\h-pgy\one_drive_prefs\OneDrive - Default Directory\Shared Documents\Estruturação do PDM 2021-2024\Elaboração PDM Versão Final')

def build_docx(filtro=None, sort_func=None, verbose=True):

    parser = ParserFichas(PATH_FILES_ONE_DRIVE/'Fichas Metas\Devolutiva 11-jun', 'rodada_5')
    fichas = parser()
    if filtro:
        print(f'Metas não encontradas {checar_metas_presentes(filtro, fichas)}')
        fichas = [ficha for ficha in fichas
                  if ficha['ficha_tecnica']['numero_meta'] in filtro]
    if sort_func:
        fichas = sort_func(fichas)
    docx = Document()
    mapbuilder = MapBuilder(path_controle=PATH_FILES_ONE_DRIVE/'Controle das Devolutivas.xlsx')
    for i, ficha in enumerate(fichas):
        if verbose:
            print(i)
        try:
            mapbuilder.create_maps(ficha)
            print(f"Ficha {ficha['file_original']} parseada com sucesso")
        except Exception as e:
            print(f"Erro ao gerar os mapas na ficha {ficha['file_original']}")
            print(e)

        builder = TableBuilder(ficha, docx)
        docx = builder()

        mapas = get_map_files(ficha['ficha_tecnica']['numero_meta'])
        for mapa in mapas:
            docx.add_picture(mapa, width = Cm(10))
            last_paragraph = docx.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        docx.add_page_break()

    docx.save('teste_10.docx')


if __name__ == '__main__':

    build_docx(filtro= None, sort_func = sort_fichas)